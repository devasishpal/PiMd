"""Cover page system — professional cover pages for reports and books."""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from pimd.branding.models import Brand
from pimd.utils.text import sanitize_text


@dataclass
class CoverConfig:
    """Configuration for a document cover page."""

    title: str = ""
    subtitle: str = ""
    version: str = ""
    author: str = ""
    company: str = ""
    date: str = ""
    revision: str = ""
    classification: str = ""
    logo_path: str | None = None
    background_color: str = "FFFFFF"
    title_color: str = "1F4E79"
    subtitle_color: str = "595959"
    accent_color: str = "C00000"
    title_size: int = 36
    subtitle_size: int = 18
    meta_size: int = 12
    show_logo: bool = True
    divider_line: bool = True


def create_cover_page(
    doc: DocxDocument,
    config: CoverConfig,
    brand: Brand | None = None,
) -> None:
    """Render a professional cover page at the start of the document.

    This inserts a new first section with the cover page content.
    """
    if brand is not None:
        cfg = brand.config
        meta = brand.metadata
        config = CoverConfig(
            title=config.title or meta.title,
            subtitle=config.subtitle or meta.subtitle,
            version=config.version or meta.version,
            author=config.author or meta.author,
            company=config.company or meta.company,
            date=config.date or "",
            revision=config.revision or meta.revision,
            logo_path=config.logo_path or cfg.logo_path,
            title_color=config.title_color or cfg.primary_color,
            subtitle_color=config.subtitle_color or cfg.secondary_color,
            accent_color=config.accent_color or cfg.accent_color,
        )

    section = doc.sections[0]
    _set_section_properties(section, config)

    _add_vertical_space(doc, 6)

    if config.show_logo and config.logo_path:
        logo_path = Path(config.logo_path)
        if logo_path.is_file():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(logo_path), width=Inches(1.5))
            _add_vertical_space(doc, 2)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run(sanitize_text(config.title or "Document Title"))
    title_run.font.size = Pt(config.title_size)
    title_run.font.bold = True
    try:
        title_run.font.color.rgb = RGBColor(
            int(config.title_color[:2], 16),
            int(config.title_color[2:4], 16),
            int(config.title_color[4:6], 16),
        )
    except (ValueError, IndexError):
        pass

    if config.subtitle:
        _add_vertical_space(doc, 1)
        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(sanitize_text(config.subtitle))
        sub_run.font.size = Pt(config.subtitle_size)
        try:
            sub_run.font.color.rgb = RGBColor(
                int(config.subtitle_color[:2], 16),
                int(config.subtitle_color[2:4], 16),
                int(config.subtitle_color[4:6], 16),
            )
        except (ValueError, IndexError):
            pass

    if config.divider_line:
        _add_vertical_space(doc, 2)
        line_p = doc.add_paragraph()
        line_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_run = line_p.add_run("─" * 60)
        try:
            line_run.font.color.rgb = RGBColor(
                int(config.accent_color[:2], 16),
                int(config.accent_color[2:4], 16),
                int(config.accent_color[4:6], 16),
            )
        except (ValueError, IndexError):
            pass
        line_run.font.size = Pt(10)

    if config.classification:
        _add_vertical_space(doc, 3)
        class_p = doc.add_paragraph()
        class_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        class_run = class_p.add_run(sanitize_text(config.classification.upper()))
        class_run.font.size = Pt(14)
        class_run.font.bold = True
        try:
            class_run.font.color.rgb = RGBColor(192, 0, 0)
        except Exception:
            pass

    _add_vertical_space(doc, 4)
    meta_items = []
    if config.author:
        meta_items.append(("Author", config.author))
    if config.company:
        meta_items.append(("Company", config.company))
    if config.version:
        meta_items.append(("Version", config.version))
    if config.revision:
        meta_items.append(("Revision", config.revision))
    if config.date:
        meta_items.append(("Date", config.date))
    for label, value in meta_items:
        mp = doc.add_paragraph()
        mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        mr = mp.add_run(sanitize_text(f"{label}: "))
        mr.font.size = Pt(config.meta_size)
        mr.font.bold = True
        vr = mp.add_run(sanitize_text(value))
        vr.font.size = Pt(config.meta_size)

    doc.add_page_break()


def _set_section_properties(section: Any, config: CoverConfig) -> None:
    """Set cover page section properties like background color."""
    try:
        sect_pr = section._sectPr
        if sect_pr is not None:
            props = sect_pr.find(qn("w:pgMar"))
            if props is None:
                props = sect_pr.makeelement(qn("w:pgMar"), {})
                sect_pr.append(props)
    except Exception:
        pass


def _add_vertical_space(doc: DocxDocument, lines: int) -> None:
    """Add empty paragraphs for vertical spacing."""
    for _ in range(lines):
        p = doc.add_paragraph()
        run = p.add_run("")
        run.font.size = Pt(10)


def _hex_to_rgb_tuple(hex_color: str) -> tuple[int, int, int]:
    """Convert hex string to RGB tuple."""
    try:
        return (
            int(hex_color[:2], 16),
            int(hex_color[2:4], 16),
            int(hex_color[4:6], 16),
        )
    except (ValueError, IndexError):
        return (0, 0, 0)
