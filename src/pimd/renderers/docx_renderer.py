"""DOCX document renderer built on ``python-docx``."""

from __future__ import annotations

import io
from datetime import date
from pathlib import Path
from typing import TYPE_CHECKING, Any

from docx import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

from pimd.exceptions import RendererError
from pimd.layout import apply_layout_to_doc
from pimd.models import (
    Block,
    Blockquote,
    BulletList,
    CodeBlock,
    Diagram,
    Document,
    EquationBlock,
    Heading,
    HorizontalRule,
    Image,
    ListItem,
    OrderedList,
    Paragraph,
    Span,
    Table,
)
from pimd.safety import SafetyGuard
from pimd.themes import ProfessionalTheme
from pimd.themes.base import Theme
from pimd.utils.logging import get_logger
from pimd.utils.text import sanitize_text

if TYPE_CHECKING:
    from docx.text.paragraph import Paragraph as DocxParagraph

try:
    from pimd.callouts import CalloutBlock as _CalloutBlock
    from pimd.callouts import callout_to_docx_element
except ImportError:
    _CalloutBlock = None  # type: ignore
    callout_to_docx_element = None  # type: ignore

logger = get_logger(__name__)

_MONO_FONT = "Courier New"
_MONO_SIZE = Pt(8.5)
_IMAGE_MAX_WIDTH = Cm(14)

# Diagram layout constants
_DIAGRAM_WIDTH_FACTOR = 0.85       # use 85% of printable width
_DIAGRAM_HEIGHT_FACTOR = 0.75      # use 75% of printable height
_DIAGRAM_MIN_WIDTH_EMU = int(3 * 914400)  # minimum 3 inches
_DIAGRAM_SPACE_BEFORE = Pt(10)
_DIAGRAM_SPACE_AFTER = Pt(8)
_DIAGRAM_CAPTION_SIZE = Pt(9)
_DIAGRAM_ERROR_SIZE = Pt(8)

# Default Markdown → DOCX style mapping
# Used when a reference document is provided.
STYLE_MAP: dict[str, str] = {
    "h1": "Heading 1",
    "h2": "Heading 2",
    "h3": "Heading 3",
    "h4": "Heading 4",
    "h5": "Heading 5",
    "h6": "Heading 6",
    "paragraph": "Normal",
    "blockquote": "Blockquote",
    "code": "Code Block",
    "table": "Table Grid",
    "caption": "Caption",
}



class DocxRenderer:
    """Render PiMD's intermediate document model as a ``.docx`` file.

    Parameters
    ----------
    theme : Theme, optional
        Visual theme that controls typography, colours, and spacing.
        Defaults to :class:`ProfessionalTheme`.
    reference_doc : str | Path | None, optional
        Path to a reference ``.docx`` file whose styles, headers, footers,
        and page layout are used as a starting point (Pandoc-style).
    style_map : dict[str, str] | None, optional
        Custom mapping of Markdown element names to DOCX style names.
        Merged over the default style map.
    """

    def __init__(
        self,
        theme: Theme | None = None,
        layout: Any | None = None,
        reference_doc: str | Path | None = None,
        style_map: dict[str, str] | None = None,
    ) -> None:
        self._doc: DocxDocument | None = None
        self._theme: Theme = theme or ProfessionalTheme()
        self._layout: Any = layout
        self._figure_counter: int = 0
        self._reference_doc_path: str | Path | None = reference_doc
        self._style_map: dict[str, str] = dict(STYLE_MAP)
        if style_map:
            self._style_map.update(style_map)

    # ------------------------------------------------------------------
    # Reference-document helpers
    # ------------------------------------------------------------------

    def _create_document(self, reference_doc: str | Path | None = None) -> DocxDocument:
        """Create a new DOCX document, optionally based on a reference doc.

        If ``reference_doc`` is provided (or was set via constructor), the
        new document inherits its styles, headers, footers, and page layout.
        Otherwise a blank document is created (default python-docx behaviour).
        """
        ref = reference_doc or self._reference_doc_path
        if ref is not None:
            ref_path = Path(ref)
            if ref_path.exists():
                logger.info("Using reference DOCX: %s", ref_path)
                return DocxDocument(str(ref_path))
            logger.warning("Reference DOCX not found, falling back to blank: %s", ref_path)
        return DocxDocument()

    def get_style_map(self) -> dict[str, str]:
        """Return the current style map (default + user overrides)."""
        return dict(self._style_map)

    def set_style_map(self, style_map: dict[str, str]) -> None:
        """Override the style map with a custom mapping."""
        self._style_map.update(style_map)

    def style_for(self, element: str) -> str | None:
        """Return the DOCX style name for a Markdown element, or ``None``.

        Checks whether the style exists in the current document before
        returning it, so callers can safely fall back.
        """
        style_name = self._style_map.get(element)
        if style_name is None:
            return None
        if self._doc is not None:
            try:
                _ = self._doc.styles[style_name]
                return style_name
            except (KeyError, AttributeError):
                pass
        return style_name

    # ------------------------------------------------------------------
    # Render
    # ------------------------------------------------------------------

    def render(
        self,
        document: Document,
        output_path: str | Path,
        *,
        generate_toc: bool = False,
        page_numbers: bool = False,
        header_text: str | None = None,
        footer_text: str | None = None,
        cover_page: bool = False,
        title: str | None = None,
        author: str | None = None,
        company: str | None = None,
        subject: str | None = None,
        keywords: list[str] | None = None,
        doc_version: str | None = None,
        reference_doc: str | Path | None = None,
        style_map: dict[str, str] | None = None,
    ) -> None:
        """Render a :class:`Document` to a ``.docx`` file.

        Args:
            document: The intermediate document model.
            output_path: Destination path for the generated file.
            generate_toc: Insert a Word TOC field at the start.
            page_numbers: Add ``Page X`` to every page footer.
            header_text: Repeating text for the page header.
            footer_text: Repeating text for the page footer.
            cover_page: Prepend a title page.
            title: Document title (used in metadata and cover page).
            author: Document author (used in metadata and cover page).
            company: Company / organisation name.
            subject: Document subject (metadata only).
            keywords: List of keywords (metadata only).
            doc_version: Version string shown on the cover page.
            reference_doc: Path to a reference ``.docx`` to use as template.
            style_map: Custom style name mappings.

        Raises:
            RendererError: If rendering fails.
        """
        try:
            if style_map:
                self._style_map.update(style_map)
            self._doc = self._create_document(reference_doc or self._reference_doc_path)
            apply_layout_to_doc(self._doc, layout=self._layout)
            self._theme.configure_styles(self._doc)
            self._set_metadata(title, author, company, subject, keywords)

            # -- Cover page (first section) --------------------------------
            if cover_page:
                self._add_cover_page(title, author, doc_version)
                self._doc.add_section()
                # Ensure headers / footers on the cover section are empty
                _clear_header_footer(self._doc.sections[0], link=False)

            # -- Table of contents -----------------------------------------
            if generate_toc:
                self._add_toc_heading()
                self._add_toc_field()
                self._doc.add_paragraph()

            # -- Main content ----------------------------------------------
            for block in document.blocks:
                self._render_block(block)

            # -- Configure headers / footers on the *last* section ---------
            content_section = self._doc.sections[-1]
            content_section.header.is_linked_to_previous = False
            content_section.footer.is_linked_to_previous = False

            if header_text:
                self._set_header(content_section, header_text)

            if page_numbers:
                self._add_page_numbers(content_section)
            elif footer_text:
                self._set_footer(content_section, footer_text)

            # Re-apply layout so that any sections added after the initial
            # call (cover, TOC, etc.) inherit the correct orientation & margins
            apply_layout_to_doc(self._doc, layout=self._layout)

            self._doc.save(str(output_path))
            logger.info("Rendered %s", output_path)
        except RendererError:
            raise
        except Exception as exc:
            raise RendererError(f"DOCX rendering failed: {exc}") from exc

    def render_to_bytes(
        self,
        document: Document,
        *,
        generate_toc: bool = False,
        page_numbers: bool = False,
        header_text: str | None = None,
        footer_text: str | None = None,
        cover_page: bool = False,
        title: str | None = None,
        author: str | None = None,
        company: str | None = None,
        subject: str | None = None,
        keywords: list[str] | None = None,
        doc_version: str | None = None,
        reference_doc: str | Path | None = None,
        style_map: dict[str, str] | None = None,
    ) -> bytes:
        """Render a :class:`Document` to DOCX bytes without writing to disk.

        Accepts the same rendering options as :meth:`render`.

        Returns:
            The DOCX file contents as ``bytes``.
        """
        try:
            if style_map:
                self._style_map.update(style_map)
            self._doc = self._create_document(reference_doc or self._reference_doc_path)
            apply_layout_to_doc(self._doc, layout=self._layout)
            self._theme.configure_styles(self._doc)
            self._set_metadata(title, author, company, subject, keywords)

            if cover_page:
                self._add_cover_page(title, author, doc_version)
                self._doc.add_section()
                _clear_header_footer(self._doc.sections[0], link=False)

            if generate_toc:
                self._add_toc_heading()
                self._add_toc_field()
                self._doc.add_paragraph()

            for block in document.blocks:
                self._render_block(block)

            content_section = self._doc.sections[-1]
            content_section.header.is_linked_to_previous = False
            content_section.footer.is_linked_to_previous = False

            if header_text:
                self._set_header(content_section, header_text)

            if page_numbers:
                self._add_page_numbers(content_section)
            elif footer_text:
                self._set_footer(content_section, footer_text)

            apply_layout_to_doc(self._doc, layout=self._layout)

            buf = io.BytesIO()
            self._doc.save(buf)
            buf.seek(0)
            logger.info("Rendered DOCX to bytes (%d bytes)", len(buf.getvalue()))
            return buf.getvalue()
        except RendererError:
            raise
        except Exception as exc:
            raise RendererError(f"DOCX rendering failed: {exc}") from exc

    # ------------------------------------------------------------------
    # Metadata
    # ------------------------------------------------------------------

    def _set_metadata(
        self,
        title: str | None,
        author: str | None,
        company: str | None,
        subject: str | None,
        keywords: list[str] | None,
    ) -> None:
        if title:
            self._doc.core_properties.title = title
        if author:
            self._doc.core_properties.author = author
        if company:
            self._doc.core_properties.category = company
        if subject:
            self._doc.core_properties.subject = subject
        if keywords:
            self._doc.core_properties.keywords = ", ".join(keywords)

    # ==================================================================
    # Cover page
    # ==================================================================

    def _add_cover_page(self, title: str | None, author: str | None, version: str | None) -> None:
        doc = self._doc

        for _ in range(6):
            doc.add_paragraph()

        # -- Title --
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sanitize_text(title or "Untitled Document"))
        run.font.size = Pt(36)
        run.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

        # -- Separator --
        self._render_horizontal_rule(doc)

        # -- Author --
        if author:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(sanitize_text(f"By {author}"))
            run.font.size = Pt(16)

        # -- Version / Date --
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        parts = []
        if version:
            parts.append(f"Version {version}")
        today = date.today().strftime("%B %d, %Y")
        parts.append(today)

        run = p.add_run(sanitize_text("  \u2022  ".join(parts)))
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

    # ==================================================================
    # Table of contents
    # ==================================================================

    def _add_toc_heading(self) -> None:
        self._doc.add_heading("Table of Contents", level=1)

    def _add_toc_field(self) -> None:
        p = self._doc.add_paragraph()

        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_begin)

        run2 = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = ' TOC \\o "1-6" \\h \\z \\u '
        run2._r.append(instr)

        run3 = p.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run3._r.append(fld_sep)

        run4 = p.add_run()
        run4.text = "[Table of Contents \u2014 right-click and select \u201cUpdate Field\u201d]"
        run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run4.italic = True
        run4.font.size = Pt(9)

        run5 = p.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run5._r.append(fld_end)

    # ==================================================================
    # Headers / Footers / Page numbers
    # ==================================================================

    @staticmethod
    def _set_header(section, text: str) -> None:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0]
        p.text = sanitize_text(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def _set_footer(section, text: str) -> None:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.text = sanitize_text(text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    @staticmethod
    def _add_page_numbers(section) -> None:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run("Page ")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        run_num = p.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        run_num._r.append(fld_begin)

        run_instr = p.add_run()
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run_instr._r.append(instr)

        run_sep = p.add_run()
        fld_sep = OxmlElement("w:fldChar")
        fld_sep.set(qn("w:fldCharType"), "separate")
        run_sep._r.append(fld_sep)

        run_placeholder = p.add_run()
        run_placeholder.text = "1"
        run_placeholder.font.size = Pt(9)
        run_placeholder.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        run_end = p.add_run()
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run_end._r.append(fld_end)

    # ------------------------------------------------------------------
    # Style setup (legacy, kept for backward compat — theme takes over)
    # ------------------------------------------------------------------

    def _configure_styles(self) -> None:
        pass  # handled by theme

    # ------------------------------------------------------------------
    # Block dispatch
    # ------------------------------------------------------------------

    def _render_block(self, block: Block) -> None:
        doc = self._doc

        if isinstance(block, Heading):
            self._render_heading(block)
        elif isinstance(block, Paragraph):
            self._render_paragraph(doc, block)
        elif isinstance(block, CodeBlock):
            self._render_code_block(doc, block)
        elif isinstance(block, Blockquote):
            self._render_blockquote(doc, block)
        elif isinstance(block, BulletList):
            self._render_bullet_list(doc, block, level=0)
        elif isinstance(block, OrderedList):
            self._render_ordered_list(doc, block, level=0)
        elif isinstance(block, Table):
            self._render_table(doc, block)
        elif isinstance(block, HorizontalRule):
            self._render_horizontal_rule(doc)
        elif isinstance(block, Image):
            self._render_image(doc, block)
        elif isinstance(block, Diagram):
            self._render_diagram(doc, block)
        elif isinstance(block, EquationBlock):
            self._render_equation_block(doc, block)
        elif isinstance(block, ListItem):
            self._render_list_item(doc, block)
        elif _CalloutBlock is not None and isinstance(block, _CalloutBlock):
            self._render_callout(doc, block)

    # ------------------------------------------------------------------
    # Heading
    # ------------------------------------------------------------------

    def _render_heading(self, block: Heading) -> None:
        level = min(max(block.level, 1), 6)
        text = sanitize_text(block.plain_text())
        p = self._doc.add_heading(text, level=level)
        if block.alignment:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = getattr(WD_ALIGN_PARAGRAPH, block.alignment.upper(), None)

    # ------------------------------------------------------------------
    # Paragraph + spans
    # ------------------------------------------------------------------

    @staticmethod
    def _render_paragraph(doc: DocxDocument, block: Paragraph) -> None:
        p = doc.add_paragraph(style="Normal")
        p.paragraph_format.space_after = Pt(6)
        if block.alignment:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = getattr(WD_ALIGN_PARAGRAPH, block.alignment.upper(), None)
        DocxRenderer._add_spans_to_paragraph(p, block.spans)

    @staticmethod
    def _add_spans_to_paragraph(p: DocxParagraph, spans: list[Span]) -> None:
        for span in spans:
            # Math span — render as PNG via PiDraw
            if span.math:
                if span.png:
                    import io

                    from docx.shared import Inches
                    try:
                        run = p.add_run()
                        run.add_picture(io.BytesIO(span.png), height=Inches(0.3))
                    except Exception:
                        run = p.add_run(f"[{span.math}]")
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                else:
                    run = p.add_run(f"[{span.math}]")
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                continue

            text = sanitize_text(span.text)
            if span.link_url:
                DocxRenderer._add_hyperlink(p, text, span.link_url)
            elif span.code:
                run = p.add_run(text)
                run.font.name = _MONO_FONT
                run.font.size = _MONO_SIZE
                run.font.color.rgb = RGBColor(0xE0, 0x3E, 0x2D)
            else:
                run = p.add_run(text)
                run.bold = span.bold
                run.italic = span.italic
                if span.underline:
                    run.underline = True
                if span.superscript:
                    DocxRenderer._apply_vert_align(run, "superscript")
                if span.subscript:
                    DocxRenderer._apply_vert_align(run, "subscript")

    @staticmethod
    def _apply_vert_align(run, kind: str) -> None:
        """Set superscript or subscript on a run via w:vertAlign."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        rpr = run._r.get_or_add_rPr()
        vert = OxmlElement("w:vertAlign")
        vert.set(qn("w:val"), kind)
        rpr.append(vert)

    @staticmethod
    def _add_hyperlink(paragraph: DocxParagraph, text: str, url: str) -> None:
        # Only external (absolute) URLs can be linked in a DOCX.
        # Relative URLs like /about or #anchor render as plain text.
        is_external = url.startswith(("http://", "https://", "mailto:", "ftp://"))
        if not is_external:
            run = paragraph.add_run(sanitize_text(text))
            run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
            run.font.underline = True
            return

        part = paragraph.part
        r_id = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        run_elem = OxmlElement("w:r")
        run_props = OxmlElement("w:rPr")

        colour = OxmlElement("w:color")
        colour.set(qn("w:val"), "0563C1")
        run_props.append(colour)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_props.append(underline)

        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")
        run_props.append(sz)

        run_elem.append(run_props)
        run_elem.text = sanitize_text(text)
        hyperlink.append(run_elem)
        paragraph._p.append(hyperlink)

    # ------------------------------------------------------------------
    # Code block
    # ------------------------------------------------------------------

    @staticmethod
    def _render_code_block(doc: DocxDocument, block: CodeBlock) -> None:
        try:
            p = doc.add_paragraph(style="Code Block")
        except KeyError:
            p = doc.add_paragraph()

        run = p.add_run(sanitize_text(block.code))
        run.font.name = _MONO_FONT
        run.font.size = _MONO_SIZE
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ------------------------------------------------------------------
    # Blockquote
    # ------------------------------------------------------------------

    def _render_blockquote(self, doc: DocxDocument, block: Blockquote) -> None:
        for child in block.children:
            if isinstance(child, Paragraph):
                try:
                    p = doc.add_paragraph(style="Blockquote")
                except KeyError:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1.0)
                DocxRenderer._add_spans_to_paragraph(p, child.spans)
            elif isinstance(child, Blockquote):
                self._render_blockquote(doc, child)
            else:
                self._render_block(child)

    # ------------------------------------------------------------------
    # Lists
    # ------------------------------------------------------------------

    def _render_bullet_list(self, doc: DocxDocument, block: BulletList, level: int) -> None:
        for item in block.items:
            self._render_list_item_with_prefix(doc, item, level, ordered=False, index=0)
            for child in item.children:
                if isinstance(child, BulletList):
                    self._render_bullet_list(doc, child, level + 1)

    def _render_ordered_list(self, doc: DocxDocument, block: OrderedList, level: int) -> None:
        for idx, item in enumerate(block.items):
            start = block.start
            self._render_list_item_with_prefix(doc, item, level, ordered=True, index=start + idx)
            for child in item.children:
                if isinstance(child, OrderedList):
                    self._render_ordered_list(doc, child, level + 1)

    def _render_list_item_with_prefix(
        self,
        doc: DocxDocument,
        item: ListItem,
        level: int,
        ordered: bool,
        index: int,
    ) -> None:
        indent = Cm(1.27 * (level + 1))
        first = True

        for child in item.children:
            if isinstance(child, Paragraph):
                p = doc.add_paragraph(style="Normal")
                p.paragraph_format.left_indent = indent
                p.paragraph_format.space_after = Pt(2)

                if first:
                    if ordered:
                        run = p.add_run(sanitize_text(f"{index}. "))
                    else:
                        bullets = ["\u2022", "\u25cb", "\u25a0"]
                        run = p.add_run(sanitize_text(f"{bullets[level % 3]} "))
                    run.font.size = Pt(10)
                    first = False

                self._add_spans_to_paragraph(p, child.spans)
            else:
                self._render_block(child)

    def _render_list_item(self, doc: DocxDocument, item: ListItem) -> None:
        for child in item.children:
            self._render_block(child)

    # ------------------------------------------------------------------
    # Table
    # ------------------------------------------------------------------

    @staticmethod
    def _render_table(doc: DocxDocument, block: Table) -> None:
        num_cols = max(len(block.headers), max((len(r) for r in block.rows), default=0))
        if num_cols == 0:
            return

        has_header = bool(block.headers)
        table = doc.add_table(rows=len(block.rows) + (1 if has_header else 0), cols=num_cols)
        table.autofit = True

        # -- Style borders via XML --
        tbl = table._tbl
        tbl_props = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        borders = OxmlElement("w:tblBorders")
        for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "333333")
            borders.append(el)
        tbl_props.append(borders)

        row_idx = 0

        if has_header:
            for col_idx, text in enumerate(block.headers):
                cell = table.rows[0].cells[col_idx]
                cell.text = sanitize_text(text)
                tc_props = cell._tc.get_or_add_tcPr()
                shading = OxmlElement("w:shd")
                shading.set(qn("w:fill"), "1A1A2E")
                shading.set(qn("w:val"), "clear")
                tc_props.append(shading)
                for par in cell.paragraphs:
                    for run in par.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.bold = True
            row_idx = 1

        for row_data in block.rows:
            for col_idx, text in enumerate(row_data):
                if col_idx < num_cols:
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = sanitize_text(text)
            row_idx += 1

        doc.add_paragraph()

    # ------------------------------------------------------------------
    # Horizontal rule
    # ------------------------------------------------------------------

    @staticmethod
    def _render_horizontal_rule(doc: DocxDocument) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        p_props = p._p.get_or_add_pPr()
        p_border = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "999999")
        p_border.append(bottom)
        p_props.append(p_border)

    # ------------------------------------------------------------------
    # Image
    # ------------------------------------------------------------------

    @staticmethod
    def _render_image(doc: DocxDocument, block: Image) -> None:
        try:
            resolved = SafetyGuard().check_path_traversal(block.url)
        except Exception:
            logger.warning("Image path blocked (security): %s", block.url)
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(sanitize_text(f"[Image: {block.alt}]"))
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return

        path = Path(resolved)
        if not path.exists():
            logger.warning("Image not found, skipping: %s", block.url)
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(sanitize_text(f"[Image: {block.alt}]"))
            run.italic = True
            run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            return

        try:
            doc.add_picture(str(path), width=_IMAGE_MAX_WIDTH)
        except Exception as exc:
            raise RendererError(f"Failed to insert image {block.url}: {exc}") from exc

    # ------------------------------------------------------------------
    # Diagram — Advanced Layout Engine
    # ------------------------------------------------------------------

    def _get_page_emu(self, doc: DocxDocument) -> tuple[int, int]:
        """Read printable area from document section in EMU."""
        section = doc.sections[-1] if doc.sections else None
        if section:
            pw = section.page_width
            ph = section.page_height
            ml = section.left_margin or Inches(1)
            mr = section.right_margin or Inches(1)
            mt = section.top_margin or Inches(1)
            mb = section.bottom_margin or Inches(1)
            return int(pw - ml - mr), int(ph - mt - mb)
        return int(6.27 * 914400), int(9.4 * 914400)

    def _get_image_size(self, png_bytes: bytes) -> tuple[int, int]:
        """Get pixel dimensions of a PNG image."""
        try:
            from PIL import Image as PILImage
            img = PILImage.open(io.BytesIO(png_bytes))
            return img.size
        except Exception:
            return 0, 0

    def _compute_diagram_dimensions(
        self, pw: int, ph: int, aspect: float
    ) -> tuple[int, int]:
        """Compute EMU dimensions from printable area and aspect ratio.

        Strategy:
        - Fit width first, using 85% of printable width
        - If resulting height exceeds 75% of printable height, constrain by height
        - Never crop — always preserve aspect ratio
        - Enforce minimum width
        """
        max_w = int(pw * _DIAGRAM_WIDTH_FACTOR)
        max_h = int(ph * _DIAGRAM_HEIGHT_FACTOR)

        # Fit by width first
        w = max_w
        h = int(w * aspect)

        # If too tall, constrain by height
        if h > max_h:
            h = max_h
            w = int(h / aspect) if aspect > 0 else max_w

        # Enforce minimum
        if w < _DIAGRAM_MIN_WIDTH_EMU:
            w = _DIAGRAM_MIN_WIDTH_EMU
            h = int(w * aspect)

        return int(w), int(h)

    def _render_diagram(self, doc: DocxDocument, block: Diagram) -> None:
        """Embed a rendered diagram with professional layout.

        Features:
        - Page-break detection: tall diagrams get pushed to next page
        - Aspect-ratio-preserving fit (never crops)
        - Transparent PNG background
        - Auto-incrementing figure numbering with bookmark
        - Professional caption formatting
        - Error placeholder when diagram data is missing
        """
        # -- Error placeholder if no image data --
        if not block.png_bytes and not block.svg_bytes:
            self._render_diagram_error(doc, block)
            return

        img_data = block.png_bytes or b""
        if not img_data and block.svg_bytes:
            logger.warning("No PNG data for diagram %s, using placeholder", block.language)
            self._render_diagram_error(doc, block)
            return

        # -- Read image dimensions and page geometry --
        pw_emu, ph_emu = self._get_page_emu(doc)
        img_pw, img_ph = self._get_image_size(img_data)

        # Validate PNG data: if PIL can't read dimensions, treat as error
        if img_pw == 0 or img_ph == 0:
            logger.warning("Invalid PNG data for diagram %s, using placeholder", block.language)
            self._render_diagram_error(doc, block)
            return

        aspect = img_ph / img_pw

        w_emu, h_emu = self._compute_diagram_dimensions(pw_emu, ph_emu, aspect)

        # -- Page-break detection: if diagram is very tall, start a new page --
        # Estimated remaining space on current page (rough heuristic)
        estimated_remaining = ph_emu * 0.4
        if h_emu > estimated_remaining and h_emu > ph_emu * 0.5:
            doc.add_page_break()

        # -- Figure numbering --
        fig_num = block.figure_number
        if fig_num is None:
            self._figure_counter += 1
            fig_num = self._figure_counter
        else:
            self._figure_counter = max(self._figure_counter, fig_num)

        # -- Embed image in centered paragraph --
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = _DIAGRAM_SPACE_BEFORE
        p.paragraph_format.space_after = Pt(2)

        try:
            img_stream = io.BytesIO(img_data)
            inline_shape = p.add_run().add_picture(
                img_stream, width=Inches(w_emu / 914400)
            )

            # Override XML extent to exact EMU (prevents Word from distorting).
            # Must also update a:ext inside pic:spPr/a:xfrm so both WP and
            # DrawingML agree — mismatched extents can cause Word to reject.
            inline = inline_shape._inline
            extent = inline.find(qn("wp:extent"))
            if extent is not None:
                extent.set("cx", str(w_emu))
                extent.set("cy", str(h_emu))

            # Sync a:xfrm/a:ext with wp:extent
            xfrm_ext = inline.findall(
                ".//" + qn("a:ext")
                + "[@cx][@cy]"
            )
            for x_ext in xfrm_ext:
                x_ext.set("cx", str(w_emu))
                x_ext.set("cy", str(h_emu))

            doc_pr = inline.find(qn("wp:docPr"))
            if doc_pr is not None:
                doc_pr.set("id", str(fig_num))
                doc_pr.set("name", f"Diagram {block.language}")
                if block.caption:
                    doc_pr.set("descr", sanitize_text(block.caption))

        except Exception as exc:
            logger.warning("Failed to embed diagram image: %s", exc)
            self._render_diagram_error(doc, block, fig_num=fig_num)
            return

        # -- Error warning below diagram (non-fatal) --
        if block.error:
            err_p = doc.add_paragraph(style="Normal")
            err_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            err_p.paragraph_format.space_before = Pt(1)
            err_p.paragraph_format.space_after = Pt(4)
            err_r = err_p.add_run(sanitize_text(f"Warning: {block.error}"))
            err_r.italic = True
            err_r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            err_r.font.size = _DIAGRAM_ERROR_SIZE

        # -- Caption with figure number --
        if block.caption:
            cap = doc.add_paragraph(style="Normal")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(4)
            cap.paragraph_format.space_after = _DIAGRAM_SPACE_AFTER

            bookmark_id = fig_num + 1000
            bookmark_name = f"fig_{fig_num}"

            bm_start = OxmlElement("w:bookmarkStart")
            bm_start.set(qn("w:id"), str(bookmark_id))
            bm_start.set(qn("w:name"), bookmark_name)
            cap._p.append(bm_start)

            caption_text = f"Figure {fig_num}: {block.caption}"
            r = cap.add_run(sanitize_text(caption_text))
            r.font.size = _DIAGRAM_CAPTION_SIZE
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            r.italic = True

            bm_end = OxmlElement("w:bookmarkEnd")
            bm_end.set(qn("w:id"), str(bookmark_id))
            cap._p.append(bm_end)

    def _render_diagram_error(
        self, doc: DocxDocument, block: Diagram, fig_num: int | None = None
    ) -> None:
        """Render a visible placeholder box when a diagram cannot be rendered.

        Shows diagram source code inside a shaded, dashed-border box
        so the reader can still see the intended content.
        """
        lang = block.language or "?"
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = _DIAGRAM_SPACE_BEFORE
        p.paragraph_format.space_after = Pt(4)

        p_props = p._p.get_or_add_pPr()
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "F5F5F5")
        p_props.append(shading)

        p_bdr = OxmlElement("w:pBdr")
        for side in ("top", "bottom", "left", "right"):
            b = OxmlElement(f"w:{side}")
            b.set(qn("w:val"), "dashed")
            b.set(qn("w:sz"), "6")
            b.set(qn("w:space"), "4")
            b.set(qn("w:color"), "AAAAAA")
            p_bdr.append(b)
        p_props.append(p_bdr)

        if block.source:
            run = p.add_run(f"[{lang} source]\n{block.source}")
        elif block.error and "not supported" in block.error.lower():
            run = p.add_run(f"[{lang} — diagram engine not available]")
        elif block.error:
            run = p.add_run(f"[{lang} — rendering failed]")
        else:
            run = p.add_run(f"[{lang} — diagram not available]")

        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run.font.size = _DIAGRAM_ERROR_SIZE
        run.italic = True

        if block.caption:
            cap = doc.add_paragraph(style="Normal")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.space_before = Pt(2)
            cap.paragraph_format.space_after = _DIAGRAM_SPACE_AFTER
            fn = fig_num or block.figure_number or 0
            caption_text = f"Figure {fn}: {block.caption}"
            r = cap.add_run(sanitize_text(caption_text))
            r.font.size = _DIAGRAM_CAPTION_SIZE
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            r.italic = True

    # ------------------------------------------------------------------
    # Callout / Admonition
    # ------------------------------------------------------------------

    @staticmethod
    def _render_callout(doc: DocxDocument, block: _CalloutBlock) -> None:  # type: ignore
        """Render a callout/admonition block as a formatted Word table.

        Delegates to ``callout_to_docx_element()`` which produces a
        ``<w:tbl>`` element with coloured left border, background fill,
        title, and content lines.
        """
        if callout_to_docx_element is None:
            p = doc.add_paragraph(style="Normal")
            run = p.add_run(sanitize_text(f"[{block.type.name}: {block.title}]"))
            run.bold = True
            for line in block.content_lines:
                lp = doc.add_paragraph(style="Normal")
                lp.add_run(sanitize_text(line))
            return

        try:
            xml_str = callout_to_docx_element(block)
            if not xml_str.strip():
                return
            tbl = OxmlElement.fromstring(xml_str.encode("utf-8"))
            doc.element.body.append(tbl)
            # Add a small spacer paragraph after the callout
            sp = doc.add_paragraph(style="Normal")
            sp.paragraph_format.space_before = Pt(2)
            sp.paragraph_format.space_after = Pt(2)
        except Exception as exc:
            logger.warning("Failed to render callout, falling back: %s", exc)
            p = doc.add_paragraph(style="Normal")
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(sanitize_text(f"[{block.type.name}: {block.title}]"))
            run.bold = True
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            for line in block.content_lines:
                lp = doc.add_paragraph(style="Normal")
                lp.add_run(sanitize_text(line))

    # ------------------------------------------------------------------
    # Equation block
    # ------------------------------------------------------------------

    @staticmethod
    def _render_equation_block(doc: DocxDocument, block: EquationBlock) -> None:
        """Render a display equation as PNG image (via PiDraw)."""
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # PNG image from PiDraw renderer
        if block.png:
            try:
                import io

                from docx.shared import Inches
                from PIL import Image as PILImage
                img = PILImage.open(io.BytesIO(block.png))
                w, h = img.size
                ar = h / max(w, 1)
                display_inches = 5.0
                run = p.add_run()
                inline_shape = run.add_picture(io.BytesIO(block.png), width=Inches(display_inches))
                # Override extent XML to match aspect ratio
                inline = inline_shape._inline
                cx = int(display_inches * 914400)
                cy = int(display_inches * 914400 * ar)
                extent = inline.find(qn("wp:extent"))
                if extent is not None:
                    extent.set("cx", str(cx))
                    extent.set("cy", str(cy))
                for x_ext in inline.findall(".//" + qn("a:ext") + "[@cx][@cy]"):
                    x_ext.set("cx", str(cx))
                    x_ext.set("cy", str(cy))
            except Exception:
                pass

        # Equation number (right-aligned)
        if block.number is not None:
            p2 = doc.add_paragraph(style="Normal")
            p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            r = p2.add_run(sanitize_text(f"({block.number})"))
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


# ======================================================================
# Helper
# ======================================================================


def _clear_header_footer(section, link: bool = False) -> None:
    """Remove all content from a section's header and footer."""
    section.header.is_linked_to_previous = link
    section.footer.is_linked_to_previous = link
    for p in section.header.paragraphs:
        p.text = ""
    for p in section.footer.paragraphs:
        p.text = ""
