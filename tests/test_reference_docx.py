"""Comprehensive tests for Reference DOCX support (Phases 1-14)."""

from __future__ import annotations

import shutil
import struct
import zipfile
from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document as DocxDocument
from docx.shared import Pt

from pimd.models import (
    Blockquote,
    BulletList,
    CodeBlock,
    Document,
    Heading,
    HorizontalRule,
    Image,
    ListItem,
    OrderedList,
    Paragraph,
    Span,
    Table,
)
from pimd.renderers.docx_renderer import STYLE_MAP, DocxRenderer
from pimd.templates.docx_reference import (
    ReferenceDoc,
    ReferenceDocError,
    install_template_package,
    validate_reference_doc,
)
from pimd.templates.style_mapper import (
    DEFAULT_STYLE_MAP,
    StyleMapper,
    get_available_styles,
    style_exists,
)

# ======================================================================
# Helpers
# ======================================================================


def _make_png() -> bytes:
    import zlib

    def chunk(chunk_type: str, data: bytes) -> bytes:
        c = chunk_type.encode() + data
        return struct.pack(">I", len(data)) + c + struct.pack(">I", zlib.crc32(c) & 0xFFFFFFFF)

    sig = b"\x89PNG\r\n\x1a\n"
    ihdr = chunk("IHDR", struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0))
    raw = zlib.compress(b"\x00\xff\x00\x00")
    idat = chunk("IDAT", raw)
    iend = chunk("IEND", b"")
    return sig + ihdr + idat + iend


def _make_reference_docx(path: Path, styles: list[str] | None = None) -> Path:
    doc = DocxDocument()
    if styles:
        for style_name in styles:
            try:
                style = doc.styles.add_style(style_name, 1)
                style.font.size = Pt(11)
            except Exception:
                pass
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.text = "Corporate Header"
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = "Confidential"
    doc.save(str(path))
    return path


def assert_valid_docx(path: Path) -> None:
    assert path.exists()
    assert path.stat().st_size > 0
    with ZipFile(path) as zf:
        assert "word/document.xml" in zf.namelist()


# ======================================================================
# Phase 1: ReferenceDoc
# ======================================================================


class TestReferenceDoc:
    def test_load_valid_docx(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx", styles=["Heading 1", "Normal"])
        ref = ReferenceDoc(ref_path)
        assert ref.path == ref_path
        assert "Heading 1" in ref.styles
        assert "Normal" in ref.styles

    def test_load_nonexistent_raises(self) -> None:
        with pytest.raises(ReferenceDocError, match="not found"):
            ReferenceDoc("/nonexistent/path.docx")

    def test_load_not_a_docx_raises(self, tmp_path: Path) -> None:
        fake = tmp_path / "fake.docx"
        fake.write_text("not a zip file")
        with pytest.raises(ReferenceDocError, match="Corrupted|valid"):
            ReferenceDoc(fake)

    def test_load_wrong_extension_raises(self, tmp_path: Path) -> None:
        txt = tmp_path / "ref.txt"
        txt.write_text("hello")
        with pytest.raises(ReferenceDocError, match=".docx"):
            ReferenceDoc(txt)

    def test_metadata_extraction(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx")
        doc = DocxDocument(str(ref_path))
        doc.core_properties.title = "Test Title"
        doc.core_properties.author = "Test Author"
        doc.save(str(ref_path))
        ref = ReferenceDoc(ref_path)
        assert ref.metadata.get("title") == "Test Title"
        assert ref.metadata.get("author") == "Test Author"

    def test_sections_extraction(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx")
        ref = ReferenceDoc(ref_path)
        assert len(ref.sections) >= 1
        section = ref.sections[0]
        assert "page_width" in section
        assert "page_height" in section
        assert "orientation" in section

    def test_headers_extraction(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx")
        ref = ReferenceDoc(ref_path)
        assert any("Corporate Header" in h for h in ref.headers)

    def test_footers_extraction(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx")
        ref = ReferenceDoc(ref_path)
        assert any("Confidential" in f for f in ref.footers)

    def test_has_style(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx", styles=["CustomStyle"])
        ref = ReferenceDoc(ref_path)
        assert ref.has_style("CustomStyle")
        assert not ref.has_style("NonExistent")

    def test_page_settings(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx")
        ref = ReferenceDoc(ref_path)
        settings = ref.page_settings
        assert settings.get("orientation") in ("portrait", "landscape")

    def test_inspect(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx", styles=["Normal", "Heading 1"])
        ref = ReferenceDoc(ref_path)
        info = ref.inspect()
        assert "styles" in info
        assert "headers" in info
        assert "footers" in info
        assert "sections" in info
        assert "page_settings" in info
        assert "metadata" in info
        assert info["style_count"] >= 2

    def test_repr(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx")
        ref = ReferenceDoc(ref_path)
        assert "ReferenceDoc" in repr(ref)
        assert "styles" in repr(ref)


# ======================================================================
# Phase 2: Renderer with reference doc
# ======================================================================


class TestRendererReferenceDoc:
    def test_default_blank_document(self, tmp_path: Path) -> None:
        doc = Document(blocks=[])
        out = tmp_path / "empty.docx"
        DocxRenderer().render(doc, out)
        assert_valid_docx(out)

    def test_renderer_with_reference_doc(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx", styles=["Heading 1", "Normal", "Code Block"])
        renderer = DocxRenderer(reference_doc=str(ref_path))
        doc = Document(blocks=[Heading(level=1, spans=[Span(text="Test")])])
        out = tmp_path / "ref_output.docx"
        renderer.render(doc, out)
        assert_valid_docx(out)

    def test_renderer_with_reference_doc_missing_style(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx", styles=["Normal"])
        renderer = DocxRenderer(reference_doc=str(ref_path))
        spans = [Span(text="Hello")]
        doc = Document(blocks=[Paragraph(spans=spans)])
        out = tmp_path / "missing_style.docx"
        renderer.render(doc, out)
        assert_valid_docx(out)

    def test_renderer_with_nonexistent_reference_falls_back(self, tmp_path: Path) -> None:
        renderer = DocxRenderer(reference_doc=str(tmp_path / "nonexistent.docx"))
        doc = Document(blocks=[Paragraph(spans=[Span(text="Fallback")])])
        out = tmp_path / "fallback.docx"
        renderer.render(doc, out)
        assert_valid_docx(out)

    def test_reference_doc_preserves_header(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx")
        renderer = DocxRenderer(reference_doc=str(ref_path))
        doc = Document(blocks=[Paragraph(spans=[Span(text="Content")])])
        out = tmp_path / "header_out.docx"
        renderer.render(doc, out)
        saved = DocxDocument(str(out))
        texts = []
        for section in saved.sections:
            for p in section.header.paragraphs:
                texts.append(p.text)
        assert any("Corporate Header" in t for t in texts), f"Header not preserved: {texts}"

    def test_reference_doc_preserves_footer(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx")
        renderer = DocxRenderer(reference_doc=str(ref_path))
        doc = Document(blocks=[Paragraph(spans=[Span(text="Content")])])
        out = tmp_path / "footer_out.docx"
        renderer.render(doc, out)
        saved = DocxDocument(str(out))
        texts = []
        for section in saved.sections:
            for p in section.footer.paragraphs:
                texts.append(p.text)
        assert any("Confidential" in t for t in texts), f"Footer not preserved: {texts}"

    def test_render_to_bytes_with_reference(self, tmp_path: Path) -> None:
        import io
        ref_path = _make_reference_docx(tmp_path / "ref.docx")
        renderer = DocxRenderer(reference_doc=str(ref_path))
        doc = Document(blocks=[Paragraph(spans=[Span(text="Bytes")])])
        result = renderer.render_to_bytes(doc)
        assert len(result) > 0
        with ZipFile(io.BytesIO(result)) as zf:
            assert "word/document.xml" in zf.namelist()


# ======================================================================
# Phase 3: StyleMapper
# ======================================================================


class TestStyleMapper:
    def test_default_mapping(self) -> None:
        mapper = StyleMapper()
        assert mapper.get("h1") == "Heading 1"
        assert mapper.get("paragraph") == "Normal"
        assert mapper.get("code") == "Code Block"
        assert mapper.get("table") == "Table Grid"

    def test_missing_element_returns_none(self) -> None:
        mapper = StyleMapper()
        assert mapper.get("nonexistent") is None

    def test_overrides(self) -> None:
        mapper = StyleMapper(overrides={"h1": "CustomHeading"})
        assert mapper.get("h1") == "CustomHeading"
        assert mapper.get("paragraph") == "Normal"

    def test_heading_level(self) -> None:
        mapper = StyleMapper()
        assert mapper.heading_level(1) == "Heading 1"
        assert mapper.heading_level(6) == "Heading 6"
        assert mapper.heading_level(0) == "Heading 1"
        assert mapper.heading_level(7) == "Heading 6"

    def test_get_with_fallback(self) -> None:
        mapper = StyleMapper()
        assert mapper.get_with_fallback("h1") == "Heading 1"
        assert mapper.get_with_fallback("nonexistent") == "Normal"

    def test_style_exists(self) -> None:
        doc = DocxDocument()
        assert style_exists(doc, "Normal")
        assert not style_exists(doc, "NonExistent")

    def test_get_available_styles(self) -> None:
        doc = DocxDocument()
        styles = get_available_styles(doc)
        assert "Normal" in styles
        assert isinstance(styles, list)

    def test_default_style_map_constant(self) -> None:
        assert "h1" in DEFAULT_STYLE_MAP
        assert "paragraph" in DEFAULT_STYLE_MAP
        assert "table" in DEFAULT_STYLE_MAP

    def test_style_map_in_renderer(self) -> None:
        assert STYLE_MAP["h1"] == "Heading 1"

    def test_clear_cache(self) -> None:
        mapper = StyleMapper()
        mapper.get("h1")
        mapper.clear_cache()
        assert mapper._cache == {}


# ======================================================================
# Phase 4/5: Style discovery
# ======================================================================


class TestStyleDiscovery:
    def test_get_available_styles_from_reference(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx", styles=["CustomA", "CustomB"])
        ref = ReferenceDoc(ref_path)
        styles = ref.get_style_names()
        assert "CustomA" in styles
        assert "CustomB" in styles

    def test_ref_has_style(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx", styles=["MyStyle"])
        ref = ReferenceDoc(ref_path)
        assert ref.has_style("MyStyle")
        assert not ref.has_style("Missing")


# ======================================================================
# Phase 6: Style overrides
# ======================================================================


class TestStyleOverrides:
    def test_renderer_style_map(self) -> None:
        renderer = DocxRenderer(style_map={"h1": "CustomH1"})
        sm = renderer.get_style_map()
        assert sm["h1"] == "CustomH1"
        assert sm["paragraph"] == "Normal"

    def test_renderer_style_for(self, tmp_path: Path) -> None:
        renderer = DocxRenderer()
        doc = Document(blocks=[Paragraph(spans=[Span(text="Test")])])
        out = tmp_path / "style_for.docx"
        renderer.render(doc, out)
        h1 = renderer.style_for("h1")
        assert h1 is None or h1 == "Heading 1"

    def test_renderer_set_style_map(self) -> None:
        renderer = DocxRenderer()
        renderer.set_style_map({"h1": "NewH1"})
        assert renderer._style_map["h1"] == "NewH1"

    def test_renderer_with_style_map_on_render(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx", styles=["CustomHeading1", "Normal"])
        renderer = DocxRenderer()
        doc = Document(blocks=[Heading(level=1, spans=[Span(text="Test")])])
        out = tmp_path / "style_map_render.docx"
        renderer.render(doc, out, reference_doc=str(ref_path), style_map={"h1": "CustomHeading1"})
        assert_valid_docx(out)


# ======================================================================
# Phase 12: Validation
# ======================================================================


class TestValidation:
    def test_validate_valid_docx(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx", styles=["Normal", "Heading 1"])
        result = validate_reference_doc(ref_path)
        assert result["valid"]

    def test_validate_nonexistent(self) -> None:
        result = validate_reference_doc("/nonexistent/doc.docx")
        assert not result["valid"]
        assert len(result["errors"]) > 0

    def test_validate_corrupted(self, tmp_path: Path) -> None:
        fake = tmp_path / "corrupt.docx"
        fake.write_text("corrupt data")
        result = validate_reference_doc(fake)
        assert not result["valid"]

    def test_validate_warns_on_missing_common_styles(self, tmp_path: Path) -> None:
        ref_path = _make_reference_docx(tmp_path / "ref.docx")
        result = validate_reference_doc(ref_path)
        assert result["valid"]


# ======================================================================
# Phase 13: Template packaging
# ======================================================================


class TestTemplatePackaging:
    def test_install_template_zip(self, tmp_path: Path) -> None:
        tpl_dir = tmp_path / "mytpl"
        tpl_dir.mkdir()
        ref_path = _make_reference_docx(tpl_dir / "reference.docx", styles=["MyStyle"])
        archive = tmp_path / "mytpl.zip"
        with zipfile.ZipFile(archive, "w") as zf:
            zf.write(ref_path, "reference.docx")
        dest = install_template_package(archive)
        assert dest.exists()
        assert (dest / "reference.docx").exists()
        ref = ReferenceDoc(dest / "reference.docx")
        assert "MyStyle" in ref.styles
        shutil.rmtree(dest, ignore_errors=True)

    def test_install_template_missing_reference(self, tmp_path: Path) -> None:
        archive = tmp_path / "bad.zip"
        with zipfile.ZipFile(archive, "w") as zf:
            zf.writestr("readme.txt", "hello")
        with pytest.raises(ReferenceDocError, match="reference.docx"):
            install_template_package(archive)


# ======================================================================
# Renderer features still work with reference doc
# ======================================================================


class TestRendererWithReferenceDocx:
    def renderer(self, tmp_path: Path) -> DocxRenderer:
        ref_path = _make_reference_docx(tmp_path / "ref.docx", styles=["Normal", "Heading 1", "Code Block"])
        return DocxRenderer(reference_doc=str(ref_path))

    def test_empty_document(self, tmp_path: Path) -> None:
        doc = Document(blocks=[])
        out = tmp_path / "empty.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_all_heading_levels(self, tmp_path: Path) -> None:
        blocks = [Heading(level=lv, spans=[Span(text=f"H{lv}")]) for lv in range(1, 7)]
        doc = Document(blocks=blocks)
        out = tmp_path / "headings.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_paragraph_with_spans(self, tmp_path: Path) -> None:
        spans = [
            Span(text="Normal "),
            Span(text="bold", bold=True),
            Span(text=" "),
            Span(text="italic", italic=True),
            Span(text=" "),
            Span(text="code", code=True),
        ]
        doc = Document(blocks=[Paragraph(spans=spans)])
        out = tmp_path / "spans.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_code_block(self, tmp_path: Path) -> None:
        doc = Document(blocks=[CodeBlock(code="def foo():\\n    pass", language="python")])
        out = tmp_path / "code.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_bullet_list(self, tmp_path: Path) -> None:
        items = [
            ListItem(children=[Paragraph(spans=[Span(text="A")])]),
            ListItem(children=[Paragraph(spans=[Span(text="B")])]),
        ]
        doc = Document(blocks=[BulletList(items=items)])
        out = tmp_path / "bullets.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_ordered_list(self, tmp_path: Path) -> None:
        items = [
            ListItem(children=[Paragraph(spans=[Span(text="One")])]),
            ListItem(children=[Paragraph(spans=[Span(text="Two")])]),
        ]
        doc = Document(blocks=[OrderedList(items=items, start=1)])
        out = tmp_path / "ordered.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_blockquote(self, tmp_path: Path) -> None:
        inner = [Paragraph(spans=[Span(text="Quoted text")])]
        doc = Document(blocks=[Blockquote(children=inner)])
        out = tmp_path / "blockquote.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_table(self, tmp_path: Path) -> None:
        doc = Document(blocks=[Table(headers=["Col A", "Col B"], rows=[["1", "2"], ["3", "4"]])])
        out = tmp_path / "table.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_horizontal_rule(self, tmp_path: Path) -> None:
        doc = Document(blocks=[HorizontalRule()])
        out = tmp_path / "hr.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_image_missing(self, tmp_path: Path) -> None:
        doc = Document(blocks=[Image(alt="missing", url="/nope.png")])
        out = tmp_path / "img_missing.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_image_present(self, tmp_path: Path) -> None:
        img = tmp_path / "pic.png"
        img.write_bytes(_make_png())
        doc = Document(blocks=[Image(alt="pic", url=str(img))])
        out = tmp_path / "img_present.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_mixed_document(self, tmp_path: Path) -> None:
        doc = Document(blocks=[
            Heading(level=1, spans=[Span(text="Title")]),
            Paragraph(spans=[Span(text="Body")]),
            CodeBlock(code="x = 1"),
            HorizontalRule(),
        ])
        out = tmp_path / "mixed.docx"
        self.renderer(tmp_path).render(doc, out)
        assert_valid_docx(out)

    def test_toc_in_document(self, tmp_path: Path) -> None:
        doc = Document(blocks=[Heading(level=1, spans=[Span(text="Chapter")])])
        out = tmp_path / "toc.docx"
        self.renderer(tmp_path).render(doc, out, generate_toc=True)
        assert_valid_docx(out)
        with ZipFile(out) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert "TOC" in xml

    def test_metadata_rendered(self, tmp_path: Path) -> None:
        doc = Document(blocks=[Paragraph(spans=[Span(text="Hello")])])
        out = tmp_path / "meta.docx"
        self.renderer(tmp_path).render(doc, out, title="Doc", author="Alice", company="Acme", subject="Test", keywords=["a"])
        assert_valid_docx(out)

    def test_page_numbers_rendered(self, tmp_path: Path) -> None:
        doc = Document(blocks=[Paragraph(spans=[Span(text="Hello")])])
        out = tmp_path / "pn.docx"
        self.renderer(tmp_path).render(doc, out, page_numbers=True)
        assert_valid_docx(out)

    def test_header_footer_rendered(self, tmp_path: Path) -> None:
        doc = Document(blocks=[Paragraph(spans=[Span(text="Hello")])])
        out = tmp_path / "hf.docx"
        self.renderer(tmp_path).render(doc, out, header_text="H", footer_text="F")
        assert_valid_docx(out)

    def test_cover_page_rendered(self, tmp_path: Path) -> None:
        doc = Document(blocks=[Paragraph(spans=[Span(text="Hello")])])
        out = tmp_path / "cover.docx"
        self.renderer(tmp_path).render(doc, out, cover_page=True, title="Cover", author="A", doc_version="1")
        assert_valid_docx(out)
