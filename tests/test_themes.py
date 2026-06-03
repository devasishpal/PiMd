"""Tests for the Theme system."""

from __future__ import annotations

import pytest
from docx import Document as DocxDocument

from pimd.themes import ProfessionalTheme, Theme


class TestThemeBase:
    def test_theme_is_abstract(self) -> None:
        """Theme base class cannot be instantiated directly."""
        with pytest.raises(TypeError):
            Theme()  # type: ignore[abstract]

    def test_theme_name_attribute(self) -> None:
        """Theme subclasses must have a name."""


class TestProfessionalTheme:
    def test_instantiation(self) -> None:
        theme = ProfessionalTheme()
        assert isinstance(theme, Theme)
        assert theme.name == "professional"

    def test_configure_styles_runs(self) -> None:
        """Applying the theme to a fresh document should not raise."""
        doc = DocxDocument()
        theme = ProfessionalTheme()
        theme.configure_styles(doc)
        # Verify styles are modified
        normal = doc.styles["Normal"]
        assert normal.font.name == "Calibri"
        assert normal.font.size is not None

    def test_configure_styles_creates_custom_styles(self) -> None:
        doc = DocxDocument()
        ProfessionalTheme().configure_styles(doc)
        # The custom styles should exist
        all_names = [s.name for s in doc.styles]
        assert "Code Block" in all_names
        assert "Blockquote" in all_names

    def test_heading_styles(self) -> None:
        doc = DocxDocument()
        ProfessionalTheme().configure_styles(doc)
        for level in range(1, 7):
            style = doc.styles[f"Heading {level}"]
            assert style.font.bold is True
            assert style.font.color.rgb is not None
